"""
Load client DOCX files into knowledge_base with embeddings (Django ORM).
Run: python manage.py load_docx [--translate] [--dry-run]
Uses config.settings (BASE_DIR, DOCX_DATA_DIR, OPENAI_API_KEY).

Document format supported:
  • Numbered Q&A blocks (primary format of the Respostas files):
        1. Question text?
        Bullet line 1
        Bullet line 2
        Sub-bullet
        2. Next question...

  • "Q:" / "Q." prefix format
  • Plain question-answer alternating paragraphs
"""
import re
import json
from pathlib import Path
from django.core.management.base import BaseCommand
from django.conf import settings
from docx import Document

RESPOSTAS_PREFIX = "Respostas "
QUESTIONS_FILE = "Lista de Perguntas - IA.docx"

# Matches lines like  "1. Question text" or "12. Something"
_NUMBERED_Q_RE = re.compile(r"^\d+\.\s+\S")


def get_paragraphs(doc_path: Path) -> list:
    """Return every non-empty paragraph text from a .docx file."""
    doc = Document(doc_path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def get_questions(doc_path: Path) -> list:
    doc = Document(doc_path)
    questions = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            questions.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in questions:
                    questions.append(t)
    return questions


# ---------------------------------------------------------------------------
# Core parser: numbered Q&A blocks (primary Respostas format)
# ---------------------------------------------------------------------------

def _is_numbered_question(text: str) -> bool:
    """Return True for lines like '1. Question text' or '12. Something?'"""
    return bool(_NUMBERED_Q_RE.match(text.strip()))


def _strip_number_prefix(text: str) -> str:
    """'1. What is X?' → 'What is X?'"""
    return re.sub(r"^\d+\.\s+", "", text.strip())


def _format_answer_lines(lines: list) -> str:
    """
    Join bullet-point lines into a single structured string.
    Lines that look like sub-headers (end with ':') are kept as-is;
    other lines get a '• ' prefix to preserve the list structure.
    """
    formatted = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Section headers like "Principais portais:" stay as-is
        if stripped.endswith(":") and len(stripped) < 80:
            formatted.append(stripped)
        else:
            formatted.append(f"• {stripped}")
    return "\n".join(formatted)


def parse_numbered_qa(paragraphs: list) -> list:
    """
    Parse Respostas .docx files whose format is:

        1. Question line
        Answer bullet 1
        Answer bullet 2
        Answer bullet 3
        2. Next question
        ...

    Returns list of (question, answer) tuples where `answer` contains ALL
    bullet-point lines that follow each numbered question, joined with newlines.
    """
    pairs = []
    current_q = None
    answer_lines = []

    for line in paragraphs:
        line = line.strip()
        if not line:
            continue

        if _is_numbered_question(line):
            # Save previous Q&A
            if current_q is not None:
                answer = _format_answer_lines(answer_lines) if answer_lines else current_q
                pairs.append((current_q, answer))
            current_q = _strip_number_prefix(line)
            answer_lines = []
        elif current_q is not None:
            # This line is part of the current answer (a bullet / sub-bullet / text)
            answer_lines.append(line)
        # Lines before the first numbered question are ignored

    # Flush the last Q&A
    if current_q is not None:
        answer = _format_answer_lines(answer_lines) if answer_lines else current_q
        pairs.append((current_q, answer))

    return pairs


# ---------------------------------------------------------------------------
# Legacy / fallback parsers (kept for other document formats)
# ---------------------------------------------------------------------------

def _starts_with_q(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    return t[0].upper() == "Q" and (len(t) == 1 or t[1] in (":", " ", "."))


def parse_qa_q_first(paragraphs: list) -> list:
    """Handle 'Q: question / A: answer' style documents."""
    pairs = []
    last_question = None
    answer_lines = []
    for line in paragraphs:
        line = line.strip()
        if not line:
            continue
        if _starts_with_q(line):
            q_text = re.sub(r"^Q\s*[:\s.]*\s*", "", line, flags=re.I).strip() or line
            if last_question is not None:
                answer_text = "\n".join(answer_lines).strip()
                pairs.append((last_question, answer_text))
            last_question = q_text
            answer_lines = []
        else:
            answer_lines.append(line)
    if last_question is not None:
        pairs.append((last_question, "\n".join(answer_lines).strip()))
    return pairs


def _is_question_line(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if t.rstrip().endswith("?"):
        return True
    if _NUMBERED_Q_RE.match(t):
        return True
    return False


def parse_question_answer_multiparagraph(paragraphs: list) -> list:
    """
    Collect multi-paragraph answers for question lines.
    Falls back if parse_numbered_qa returns nothing.
    """
    pairs = []
    i = 0
    while i < len(paragraphs):
        line = paragraphs[i].strip()
        if not line:
            i += 1
            continue
        if not _is_question_line(line):
            i += 1
            continue
        q = _strip_number_prefix(line) if _is_numbered_question(line) else line
        answer_lines = []
        j = i + 1
        while j < len(paragraphs):
            next_line = paragraphs[j].strip()
            if not next_line:
                j += 1
                continue
            if _is_question_line(next_line):
                break
            answer_lines.append(next_line)
            j += 1
        full_answer = _format_answer_lines(answer_lines) if answer_lines else q
        pairs.append((q, full_answer))
        i = j if answer_lines else i + 1
    return pairs


def parse_qa_pairs(paragraphs: list) -> list:
    """
    Handle 'Pergunta: ...' / 'P: ...' explicit prefix format.
    Does NOT match plain numbered lines (those go through parse_numbered_qa).
    """
    pairs = []
    i = 0
    while i < len(paragraphs):
        line = paragraphs[i]
        q_match = re.match(r"^(?:P|Pergunta|Q|Question)\s*[:\d.]*\s*(.+)", line, re.I)
        if q_match and i + 1 < len(paragraphs):
            pairs.append((q_match.group(1).strip(), paragraphs[i + 1].strip()))
            i += 2
            continue
        i += 1
    return pairs


def parse_question_answer_paragraphs(paragraphs: list) -> list:
    """Simple alternating Q/A fallback."""
    pairs = []
    i = 0
    while i < len(paragraphs):
        q = paragraphs[i].strip()
        if not q:
            i += 1
            continue
        if q.rstrip().endswith("?") and i + 1 < len(paragraphs):
            a = paragraphs[i + 1].strip()
            if a:
                pairs.append((q, a))
                i += 2
                continue
        pairs.append((q, q))
        i += 1
    return pairs


def _choose_best_parser(paragraphs: list, questions: list = None) -> list:
    """
    Select the best parser for a set of paragraphs.
    Priority order:
      1. Numbered Q&A blocks  (1. Question / bullet1 / bullet2 …)  ← primary Respostas format
      2. Q-prefixed blocks    (Q: / Q. / Q  format)
      3. Exact question↔paragraph alignment  (one paragraph per question)
      4. Multi-paragraph fallback
      5. Explicit P:/Pergunta: prefix
      6. Simple alternating Q/A
      7. Last resort: each paragraph is its own entry
    """
    # 1. Numbered Q&A (Respostas files: "1. Question\nbullet\nbullet\n2. ...")
    pairs = parse_numbered_qa(paragraphs)
    if pairs:
        return pairs

    # 2. Q-prefix style  ("Q: something / answer / Q: next...")
    pairs = parse_qa_q_first(paragraphs)
    if pairs:
        return pairs

    # 3. One-to-one alignment with questions file
    if questions and len(paragraphs) == len(questions):
        return list(zip(questions, paragraphs))

    # 4. Multi-paragraph (any line ending '?' or starting with digit+dot is a question)
    pairs = parse_question_answer_multiparagraph(paragraphs)
    if pairs:
        return pairs

    # 5. Explicit prefix  (P:/Pergunta:)
    pairs = parse_qa_pairs(paragraphs)
    if pairs:
        return pairs

    # 6. Alternating Q/A
    pairs = parse_question_answer_paragraphs(paragraphs)
    if pairs:
        return pairs

    # 7. Last resort
    return [(p, p) for p in paragraphs]


def translate_to_english(text: str) -> str:
    if not text or not text.strip():
        return text or ""
    if not getattr(settings, "OPENAI_API_KEY", None):
        return text
    try:
        from openai import OpenAI
        client = OpenAI(api_key=settings.OPENAI_API_KEY)
        r = client.chat.completions.create(
            model=getattr(settings, "GPT_MODEL", "gpt-4o-mini"),
            messages=[
                {"role": "system", "content": "Translate the following text to English. Preserve meaning and structure. Output only the translation, no bullets or numbering unless in the original."},
                {"role": "user", "content": text[:6000]},
            ],
            temperature=0.2,
        )
        out = (r.choices[0].message.content or "").strip()
        return out if out else text
    except Exception:
        return text


class Command(BaseCommand):
    help = "Load client DOCX files into knowledge_base with embeddings (state/county, document_source). Optionally translate answers to English."

    def add_arguments(self, parser):
        parser.add_argument("--translate", action="store_true", help="Translate Spanish/Portuguese answers to English before storing")
        parser.add_argument("--dry-run", action="store_true", help="Show what would be loaded without saving")

    def handle(self, *args, **options):
        from chatbot.models import KnowledgeBase
        from chatbot.services.knowledge_service import get_embedding, build_kb_search_fields

        translate = options.get("translate", False)
        dry_run = options.get("dry_run", False)
        if dry_run:
            self.stdout.write("DRY RUN: no changes will be saved.")

        _backend_dir = Path(settings.BASE_DIR)
        _parent = _backend_dir.parent
        docx_dir = getattr(settings, "DOCX_DATA_DIR", None)
        if docx_dir is not None:
            docx_dir = Path(docx_dir) if not isinstance(docx_dir, Path) else docx_dir
        candidates = [
            docx_dir,
            _parent / "documents",
            _parent / "data",
            _parent / "chatbot_documents",
            _parent,
            _backend_dir / "chatbot" / "data",
            _backend_dir / "data",
            _backend_dir / "documents",
            _backend_dir,
        ]
        data_dir = None
        for d in candidates:
            if d is None:
                continue
            if isinstance(d, str):
                d = Path(d)
            if d.is_dir():
                if (d / QUESTIONS_FILE).exists() or list(d.glob("Respostas *.docx")):
                    data_dir = d
                    break
        if data_dir is None:
            data_dir = docx_dir if (docx_dir and docx_dir.is_dir()) else (_parent / "documents" if (_parent / "documents").is_dir() else _backend_dir)
        self.stdout.write(f"Using data dir: {data_dir}")

        questions_file = data_dir / QUESTIONS_FILE
        questions = []
        if questions_file.exists():
            questions = get_questions(questions_file)
            self.stdout.write(f"Loaded {len(questions)} questions from {QUESTIONS_FILE}")
        else:
            self.stdout.write(f"Questions file not found: {questions_file}")

        state_answers = {}
        for path in sorted(data_dir.glob("*.docx")):
            name = path.stem
            if name.startswith(RESPOSTAS_PREFIX):
                state = name[len(RESPOSTAS_PREFIX):].strip()
                paras = get_paragraphs(path)
                if not paras:
                    continue
                pairs = _choose_best_parser(paras, questions)
                state_answers[state] = pairs
                self.stdout.write(f"  {state}: {len(state_answers[state])} entries")

        if not state_answers and not questions:
            self.stdout.write("No DOCX data found. Set DOCX_DATA_DIR in .env to your folder (can be outside braelo), or use: braelo_backend/documents/, braelo/chatbot/data/, etc.")
            self.stdout.write("Expected: Lista de Perguntas - IA.docx, Respostas Arizona.docx, Respostas Texas.docx, etc.")
            return

        entries = []
        for state, qa_list in state_answers.items():
            for q, a in qa_list:
                if q and a:
                    entries.append((state, q, a))
        for q in questions:
            if q and not any(e[1] == q for e in entries):
                entries.append((None, q, "See state-specific answers or ask for your state."))

        if not getattr(settings, "OPENAI_API_KEY", None):
            self.stdout.write("OPENAI_API_KEY not set. Embeddings will be empty; semantic search will not work until you add key and re-run.")

        if not dry_run:
            existing = KnowledgeBase.objects.count()
            if existing > 0:
                KnowledgeBase.objects.all().delete()
                self.stdout.write(f"Cleared {existing} existing knowledge_base rows.")

        added = 0
        for state, question, answer in entries:
            if translate:
                answer = translate_to_english(answer)
                if added == 0:
                    self.stdout.write("Translating answers to English...")
            if dry_run:
                added += 1
                if added <= 3:
                    self.stdout.write(f"  [dry-run] state={state}, q={question[:50]}...")
                continue
            answer_snippet = (answer or "")[:300].replace("\n", " ").strip()
            embed_text = question if not answer_snippet else f"{question} {answer_snippet}"
            emb = get_embedding(embed_text)
            emb_json = json.dumps(emb) if emb else None
            idx = build_kb_search_fields(question, answer)
            KnowledgeBase.objects.create(
                state=state,
                county=None,
                question=question,
                answer=answer,
                embedding_json=emb_json,
                search_tokens=idx.get("search_tokens"),
                document_source=state and f"Respostas {state}.docx" or "Lista de Perguntas - IA.docx",
            )
            added += 1
            if added % 10 == 0:
                self.stdout.write(f"  Added {added}/{len(entries)}...")

        if dry_run:
            self.stdout.write(self.style.WARNING(f"Dry run: would insert {added} rows."))
        else:
            self.stdout.write(self.style.SUCCESS(f"Done. Inserted {added} rows into knowledge_base."))
